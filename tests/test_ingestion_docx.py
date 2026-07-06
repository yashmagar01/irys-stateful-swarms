"""Unit tests for docx ingestion (S-14).

Builds a minimal .docx with python-docx, then verifies that read_docx()
extracts paragraph and table text correctly via the fallback path.
The pandoc path is also exercised if pandoc is available on PATH.
"""
from __future__ import annotations

import shutil
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.shared import Pt

from src.ingestion.docx import read_docx


@pytest.fixture()
def simple_docx(tmp_path: Path) -> Path:
    """Three paragraphs: heading, body, footer."""
    doc = DocxDocument()
    doc.add_heading("Report Title", level=1)
    doc.add_paragraph("This is the first body paragraph.")
    doc.add_paragraph("This is the second body paragraph.")
    p = tmp_path / "report.docx"
    doc.save(str(p))
    return p


@pytest.fixture()
def docx_with_table(tmp_path: Path) -> Path:
    """A doc with a paragraph and a 2x3 table."""
    doc = DocxDocument()
    doc.add_paragraph("Intro text")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Col A"
    table.cell(0, 1).text = "Col B"
    table.cell(0, 2).text = "Col C"
    table.cell(1, 0).text = "Val 1"
    table.cell(1, 1).text = "Val 2"
    table.cell(1, 2).text = "Val 3"
    p = tmp_path / "table.docx"
    doc.save(str(p))
    return p


class TestReadDocxParagraphs:
    def test_heading_extracted(self, simple_docx):
        text, _ = read_docx(simple_docx)
        assert "Report Title" in text

    def test_body_paragraphs_extracted(self, simple_docx):
        text, _ = read_docx(simple_docx)
        assert "first body paragraph" in text
        assert "second body paragraph" in text

    def test_returns_string_and_dict(self, simple_docx):
        text, meta = read_docx(simple_docx)
        assert isinstance(text, str)
        assert isinstance(meta, dict)


class TestReadDocxTable:
    def test_table_cells_extracted(self, docx_with_table):
        text, _ = read_docx(docx_with_table)
        assert "Col A" in text
        assert "Val 1" in text

    def test_paragraph_and_table_both_present(self, docx_with_table):
        text, _ = read_docx(docx_with_table)
        assert "Intro text" in text
        assert "Col B" in text


class TestReadDocxErrorHandling:
    def test_nonexistent_file_returns_error_string(self, tmp_path):
        bad = tmp_path / "missing.docx"
        text, meta = read_docx(bad)
        assert "error" in text.lower()
        assert meta == {}
