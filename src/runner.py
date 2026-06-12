from __future__ import annotations

import json
import os
import re
import time
from dataclasses import asdict, dataclass, field
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt

from .ingestion import discover_documents
from .providers.gemini import GeminiCaller
from .swarm import run_swarm
from .swarm.models import Task
from .swarm.survival_trace import (
    extract_artifact_texts,
    finalize_survival_trace,
)


@dataclass
class RunResult:
    task_id: str
    score: float | None = None
    criteria_pass: int = 0
    criteria_total: int = 0
    tokens_used: int = 0
    cost_usd: float = 0.0
    wall_clock_seconds: float = 0.0
    deliverable_files: list[str] = field(default_factory=list)
    error: str | None = None


_GENERATION_METADATA_BLOCKLIST = frozenset({
    "task_id", "criteria", "match_criteria", "scorer", "scores",
    "criteria_results", "rubric", "prior_score", "prior_scores",
})


def _generation_metadata(task_data: dict, deliverables: dict) -> dict:
    """Build metadata safe for generation — no evaluator/benchmark fields."""
    return {
        "title": task_data.get("title", ""),
        "work_type": task_data.get("work_type", ""),
        "deliverables": deliverables,
    }


def run_single_task(task_dir: Path, output_dir: Path, *,
                    worker_model: str | None = None,
                    synthesis_model: str | None = None,
                    task_id: str | None = None) -> RunResult:
    if not task_id:
        task_id = f"{task_dir.parent.name}/{task_dir.name}"
    t0 = time.time()

    task_json_path = task_dir / "task.json"
    if not task_json_path.exists():
        return RunResult(task_id=task_id, error="task.json not found")

    task_data = json.loads(task_json_path.read_text(encoding="utf-8-sig"))
    deliverables_for_task = _deliverables_for_task(task_data)

    instruction = task_data.get("instructions", "")
    if not instruction:
        instruction_path = task_dir / "instruction.md"
        if instruction_path.exists():
            instruction = instruction_path.read_text(encoding="utf-8")

    if not instruction:
        return RunResult(task_id=task_id, error="no instruction found")

    documents = discover_documents(task_dir, task_data.get("docs_path"))
    if not documents:
        return RunResult(task_id=task_id, error="no documents found")

    w_model = worker_model or os.getenv("SWARM_WORKER_MODEL", "gemini-3.1-flash-lite")
    s_model = synthesis_model or os.getenv("SWARM_SYNTHESIS_MODEL", "gemini-3.5-flash")

    r_model = os.getenv("SWARM_REVIEWER_MODEL", "gemini-3.5-flash")

    worker_caller = GeminiCaller(model=w_model)
    synthesis_caller = GeminiCaller(model=s_model) if s_model != w_model else worker_caller
    reviewer_caller = GeminiCaller(model=r_model) if r_model else None

    f_model = os.getenv("SWARM_FABLE_MODEL", "")
    if f_model and reviewer_caller is not None:
        from .providers.anthropic import AnthropicCaller
        from .providers.rotating import RotatingCaller
        fable_caller = AnthropicCaller(model=f_model)
        reviewer_caller = RotatingCaller(
            [fable_caller, reviewer_caller], pattern=[0, 1, 1],
        )

    out_dir = output_dir / task_id.replace("/", os.sep)
    out_dir.mkdir(parents=True, exist_ok=True)
    output_subdir = out_dir / "output"
    output_subdir.mkdir(exist_ok=True)

    task = Task(
        instruction=instruction,
        documents=documents,
        metadata=_generation_metadata(task_data, deliverables_for_task),
        output_dir=str(out_dir),
    )

    try:
        deliverable, blackboard = run_swarm(
            task, worker_caller,
            synthesis_caller=synthesis_caller,
            reviewer_caller=reviewer_caller,
        )
    except Exception as e:
        return RunResult(task_id=task_id, error=f"swarm error: {e}")

    deliverable_files = _write_deliverables(
        deliverable, deliverables_for_task, output_subdir,
    )
    artifact_texts = extract_artifact_texts(output_subdir, deliverable_files)
    finalize_survival_trace(out_dir, artifact_texts)

    wall_clock = time.time() - t0
    tokens = blackboard.total_tokens_used
    tokens_in = blackboard.tokens_input
    tokens_out = blackboard.tokens_output

    # Cost calculation per model
    MODEL_PRICING = {
        "gemini-3.1-flash-lite": {"input": 0.25, "output": 1.50},
        "gemini-3-flash-preview": {"input": 0.50, "output": 3.00},
        "gemini-3.5-flash": {"input": 1.50, "output": 9.00},
        "gemini-3.1-pro-preview": {"input": 2.00, "output": 12.00},
        "claude-fable-5": {"input": 10.00, "output": 50.00},
    }
    DEFAULT_PRICING = {"input": 0.25, "output": 1.50}

    cost_total = 0.0
    cost_breakdown = {}
    for model, usage in blackboard.cost_by_model.items():
        pricing = MODEL_PRICING.get(model, DEFAULT_PRICING)
        model_cost_in = usage["input"] * pricing["input"] / 1_000_000
        model_cost_out = usage["output"] * pricing["output"] / 1_000_000
        model_cost = model_cost_in + model_cost_out
        cost_total += model_cost
        cost_breakdown[model] = {
            "input_tokens": usage["input"],
            "output_tokens": usage["output"],
            "total_tokens": usage["total"],
            "calls": usage["calls"],
            "cost_usd": round(model_cost, 4),
        }

    # Fallback if no per-model tracking
    if not cost_breakdown:
        fallback_pricing = DEFAULT_PRICING
        cost_total = (tokens_in * fallback_pricing["input"] + tokens_out * fallback_pricing["output"]) / 1_000_000

    metrics = {
        "documents_read": len(documents),
        "total_vdr_files": len(documents),
        "documents_skipped": 0,
        "documents_read_list": [d.name for d in documents],
        "documents_skipped_list": [],
        "input_tokens": tokens_in,
        "output_tokens": tokens_out,
        "total_tokens": tokens,
        "cost_total_usd": round(cost_total, 4),
        "cost_by_model": cost_breakdown,
        "wall_clock_seconds": wall_clock,
    }
    (out_dir / "metrics.json").write_text(
        json.dumps(metrics, indent=2), encoding="utf-8",
    )

    status = {
        "task_id": task_id,
        "status": "completed",
        "tokens_used": tokens,
        "wall_clock_seconds": wall_clock,
        "deliverable_files": deliverable_files,
    }
    (out_dir / "status.json").write_text(
        json.dumps(status, indent=2), encoding="utf-8",
    )

    return RunResult(
        task_id=task_id,
        tokens_used=tokens,
        wall_clock_seconds=wall_clock,
        deliverable_files=deliverable_files,
    )


def _write_deliverables(deliverable: str | dict[str, str], deliverables_map: dict,
                        output_dir: Path) -> list[str]:
    files_written = []

    if deliverables_map:
        for filename in deliverables_map.values():
            path = output_dir / filename
            suffix = Path(filename).suffix.lower()
            content = _content_for_file(deliverable, filename)
            if suffix == ".docx":
                _write_docx(path, content)
            elif suffix == ".xlsx":
                _write_xlsx(path, content)
            else:
                path.write_text(content, encoding="utf-8")
            files_written.append(filename)
    else:
        _write_docx(output_dir / "output.docx", _content_for_file(deliverable, "output.docx"))
        files_written.append("output.docx")

    return files_written


def _deliverables_for_task(task_data: dict) -> dict[str, str]:
    deliverables_map = task_data.get("deliverables", {})
    if isinstance(deliverables_map, dict) and deliverables_map:
        return {
            str(key): str(value)
            for key, value in deliverables_map.items()
            if isinstance(value, str) and value.strip()
        }
    return {}


def _content_for_file(deliverable: str | dict[str, str], filename: str) -> str:
    if isinstance(deliverable, dict):
        if filename in deliverable:
            return deliverable[filename]
        stem = Path(filename).name
        if stem in deliverable:
            return deliverable[stem]
        available = ", ".join(sorted(str(k) for k in deliverable)) or "<none>"
        raise KeyError(
            f"No synthesized content for {filename}; available files: {available}"
        )
    return deliverable


def _write_docx(path: Path, text: str) -> None:
    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r'^\d+[.)] ', stripped):
            text_after = re.sub(r'^\d+[.)] ', '', stripped)
            doc.add_paragraph(text_after, style="List Number")
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        else:
            doc.add_paragraph(stripped)

    doc.save(str(path))


def _write_xlsx(path: Path, text: str) -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Analysis"

    def sanitize_sheet_name(name: str) -> str:
        cleaned = re.sub(r"[\[\]\*\?/\\:]", " ", name).strip() or "Sheet"
        return cleaned[:31]

    def ensure_sheet(name: str):
        title = sanitize_sheet_name(name)
        if title in wb.sheetnames:
            return wb[title]
        if not ws._cells and ws.title == "Analysis":
            ws.title = title
            return ws
        return wb.create_sheet(title)

    current = ws
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        sheet_match = re.match(r"^#{1,3}\s*Sheet:\s*(.+)$", stripped, re.IGNORECASE)
        if sheet_match:
            current = ensure_sheet(sheet_match.group(1))
            continue
        if re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", stripped):
            continue
        if "|" in stripped and stripped.count("|") >= 2:
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            current.append(cells)
        else:
            current.append([stripped])

    wb.save(str(path))
