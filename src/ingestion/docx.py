from __future__ import annotations

import shutil
import subprocess
from pathlib import Path


def read_docx(path: Path) -> tuple[str, dict]:
    pandoc = shutil.which("pandoc")
    if not pandoc:
        return _read_docx_fallback(path), {}

    result = subprocess.run(
        [pandoc, str(path), "-t", "markdown", "--wrap=none", "--track-changes=accept"],
        capture_output=True, text=True, timeout=30,
        encoding="utf-8", errors="replace",
    )
    if result.returncode != 0:
        return _read_docx_fallback(path), {}
    return result.stdout, {}


def _read_docx_fallback(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        parts = []
        for para in doc.paragraphs:
            parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                parts.append("\t".join(cells))
        return "\n".join(parts)
    except Exception as e:
        return f"(error reading {path.name}: {e})"
